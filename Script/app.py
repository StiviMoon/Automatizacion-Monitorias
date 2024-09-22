import openpyxl
from openpyxl import load_workbook
from datetime import datetime, timedelta
from docx import Document
import os


# Cargar archivo Excel
def load_data_from_excel(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"El archivo {file_path} no se encontró.")

    try:
        workbook = load_workbook(filename=file_path)
    except Exception as e:
        raise Exception(f"Error al cargar el archivo de Excel: {e}")

    sheet = workbook.active

    # Leer encabezados (opcional) y limpiar espacios en blanco
    headers = [cell.value.strip() for cell in sheet[1]]  # Limpia espacios en blanco

    # Leer filas de datos (comienza desde la fila 2)
    data = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        data.append(dict(zip(headers, row)))

    return data


# Filtrar monitorías de la semana actual
def filter_monitorias_week(data):
    today = datetime.now()
    start_week = today - timedelta(days=today.weekday())  # Inicio de semana (lunes)
    end_week = start_week + timedelta(days=6)  # Fin de semana (domingo)

    monitorias_week = []
    for entry in data:
        if "Fecha de la monitoría" in entry:
            fecha_monitoria = entry["Fecha de la monitoría"]

            # Verificar si es un string y convertir, si es necesario
            if isinstance(fecha_monitoria, str):
                try:
                    fecha_monitoria = datetime.strptime(fecha_monitoria, "%d/%m/%Y")
                except ValueError:
                    print(f"Formato de fecha incorrecto en: {entry}")
                    continue  # Saltar esta entrada si la fecha es incorrecta

            # Si la fecha ya es un objeto datetime, continuar
            if (
                isinstance(fecha_monitoria, datetime)
                and start_week <= fecha_monitoria <= end_week
            ):
                monitorias_week.append(entry)

    return monitorias_week


# Crear resumen semanal
def create_weekly_summary(monitorias_week):
    document = Document()
    document.add_heading(f"Resumen Semanal de Monitorías", 0)

    today = datetime.now()
    week_range = f"Semana del {today - timedelta(days=today.weekday()):%d/%m/%Y} al {(today + timedelta(days=6 - today.weekday())):%d/%m/%Y}"
    document.add_paragraph(f"Período: {week_range}")

    if monitorias_week:
        for monitoria in monitorias_week:

            document.add_paragraph("----------")
            document.add_paragraph(
                f"Estudiante: {monitoria['Nombre completo del estudiante']}"
            )
            document.add_paragraph(f"Código: {monitoria['Código del estudiante']}")
            document.add_paragraph(
                f"Correo: {monitoria['Dirección de correo electrónico']}"
            )
            document.add_paragraph(f"Grupo Académico: {monitoria['Grupo académico']}")
            document.add_paragraph(f"Jornada: {monitoria['Jornada de estudios']}")
            document.add_paragraph(
                f"Tipo de Monitoría: {monitoria['Tipo de monitoría recibida']}"
            )
            document.add_paragraph(
                f"Fecha de la Monitoría: {monitoria['Fecha de la monitoría']}"
            )
            document.add_paragraph(f"Horario: {monitoria['Horario de la monitoría']}")
            document.add_paragraph(
                f"Modalidad: {monitoria['Modalidad de la monitoría']}"
            )
            document.add_paragraph(
                f"Comentarios: {monitoria['Comentarios adicionales']}"
            )
            document.add_paragraph("----------")

    else:
        document.add_paragraph("----------")
        document.add_paragraph("No hubieron agendaciones esta semana.")
        document.add_paragraph(
            "Se trabajó en la creación de materiales de apoyo para los estudiantes."
        )
        document.add_paragraph("----------")

    return document


# Guardar resumen en un documento Word
def save_summary(document, save_path):
    # Sobrescribimos el archivo si ya existe
    document.save(save_path)
    print(f"Resumen semanal guardado en: {save_path}")


# Guardar log de monitorías evitando duplicados
def save_log(data, log_file):
    existing_data = set()

    # Si el archivo de bitácora existe, cargar datos ya registrados
    if os.path.exists(log_file):
        workbook = load_workbook(log_file)
        sheet = workbook.active

        # Leer datos existentes y almacenarlos como tuplas (para evitar duplicados)
        for row in sheet.iter_rows(min_row=2, values_only=True):
            existing_data.add(tuple(row))
    else:
        # Si no existe, crear un nuevo archivo
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.append(list(data[0].keys()))  # Encabezados

    # Añadir nuevas entradas (si no están ya en la bitácora)
    new_entries_count = 0
    for entry in data:
        entry_tuple = tuple(entry.values())
        if entry_tuple not in existing_data:
            sheet.append(list(entry.values()))
            new_entries_count += 1

    workbook.save(log_file)
    print(f"Bitácora actualizada en: {log_file}. Nuevas entradas: {new_entries_count}")


# Mostrar el directorio de trabajo actual
print(f"Directorio de trabajo actual: {os.getcwd()}")

# Actualiza la ruta del archivo Excel cargado
excel_file_path = os.path.abspath("./Datos-excel/Formulario-FPI.xlsx")

# Ruta para guardar el resumen semanal y la bitácora
summary_file_path = os.path.abspath("./Recopilacion/resumen_semanal1.docx")
log_file_path = os.path.abspath("./Recopilacion/bitacora_monitorias.xlsx")

try:
    # Cargar los datos desde el archivo Excel
    data = load_data_from_excel(excel_file_path)

    # Filtrar monitorías de la semana actual
    monitorias_week = filter_monitorias_week(data)

    # Crear y guardar el resumen semanal
    document = create_weekly_summary(monitorias_week)
    save_summary(document, summary_file_path)

    # Guardar la bitácora actualizada
    save_log(data, log_file_path)

except Exception as e:
    print(f"Ocurrió un error: {e}")
